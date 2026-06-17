from datetime import datetime
import os
from docx import Document
from docx.shared import Inches


class ScreenshotManager:

    def __init__(self):
        self.counter = 1
        self.screenshot_paths = []
        self.test_case_name = None
        self.screenshot_dir = None

    def start_test(self, test_case_name):
        """
        Initialize screenshot folder for a test case.
        """
        self.test_case_name = test_case_name

        self.screenshot_dir = os.path.join(
            "reports",
            test_case_name,
            "screenshots"
        )

        os.makedirs(self.screenshot_dir, exist_ok=True)

    def capture(self, page, step_name):
        """
        Capture screenshot and store path with timestamp.
        """
        if not self.screenshot_dir:
            raise Exception(
                "start_test() must be called before capture()."
            )

        # Wait until page is rendered
        page.wait_for_load_state("domcontentloaded")

        try:
            page.wait_for_load_state(
                "networkidle",
                timeout=5000
            )
        except Exception:
            pass

        page.wait_for_timeout(1500)

        timestamp = datetime.now().strftime(
            "%Y-%m-%d %H:%M:%S"
        )

        timestamp_for_file = datetime.now().strftime(
            "%Y-%m-%d_%H-%M-%S"
        )

        safe_step_name = step_name.replace(
            " ",
            "_"
        )

        file_name = (
            f"{self.counter:02d}_{safe_step_name}_"
            f"{timestamp_for_file}.png"
        )

        file_path = os.path.join(
            self.screenshot_dir,
            file_name
        )

        page.screenshot(
            path=file_path,
            full_page=True
        )

        self.screenshot_paths.append(
            (
                self.counter,
                step_name,
                timestamp,
                file_path
            )
        )

        self.counter += 1

    def create_word_report(self, test_case_name=None):
        """
        Create Word report with all screenshots and timestamps.
        """
        if test_case_name:
            self.test_case_name = test_case_name

        if not self.test_case_name:
            raise Exception(
                "Test case name not provided."
            )

        report_dir = os.path.join(
            "reports",
            self.test_case_name
        )

        os.makedirs(
            report_dir,
            exist_ok=True
        )

        report_path = os.path.join(
            report_dir,
            f"{self.test_case_name}.docx"
        )

        doc = Document()

        doc.add_heading(
            f"Execution Report - {self.test_case_name}",
            level=1
        )

        for (
            step_no,
            step_name,
            timestamp,
            image_path
        ) in self.screenshot_paths:

            if not os.path.exists(
                image_path
            ):
                continue

            doc.add_heading(
                f"Step {step_no}: {step_name}",
                level=2
            )

            doc.add_paragraph(
                f"Timestamp: {timestamp}"
            )

            doc.add_picture(
                image_path,
                width=Inches(5.5)
            )

        doc.save(
            report_path
        )

        print(
            f"Word report generated: {report_path}"
        )

    def capture_failure(
        self,
        page,
        error_message="Test Failed"
    ):
        """
        Capture screenshot when test fails.
        """
        if not self.screenshot_dir:
            raise Exception(
                "start_test() must be called before capture_failure()."
            )

        page.wait_for_load_state(
            "domcontentloaded"
        )

        page.wait_for_timeout(
            1000
        )

        timestamp = datetime.now().strftime(
            "%Y-%m-%d %H:%M:%S"
        )

        timestamp_for_file = datetime.now().strftime(
            "%Y-%m-%d_%H-%M-%S"
        )

        file_name = (
            f"{self.counter:02d}_FAILED_"
            f"{timestamp_for_file}.png"
        )

        file_path = os.path.join(
            self.screenshot_dir,
            file_name
        )

        page.screenshot(
            path=file_path,
            full_page=True
        )

        self.screenshot_paths.append(
            (
                self.counter,
                f"FAILED - {error_message}",
                timestamp,
                file_path
            )
        )

        self.counter += 1